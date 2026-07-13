from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from dataclasses import dataclass


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

@dataclass
class LoadedDocument:
    filename: str
    text: str
    page: int | None = None


def load_document(file_path: Path) -> LoadedDocument:
    """
    Load a supported document and return its text.

    Args:
        file_path: Path to the document.

    Returns:
        Extracted text.

    Raises:
        ValueError: If the file extension is not supported.
    """

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(file_path)

    if suffix == ".docx":
        return _load_docx(file_path)

    raise ValueError(f"Unsupported file type: {suffix}")


def _load_docx(file_path: Path) -> LoadedDocument:
    """Extract text from a Word document."""

    document = Document(file_path)

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    return LoadedDocument(filename=file_path.name, text="\n".join(paragraphs))


def _load_pdf(file_path: Path):
    source = Path(file_path).name
    pages = []

    with fitz.open(file_path) as pdf:

        for page_number, page in enumerate(pdf, start=1):

            text = page.get_text()

            if text.strip():

                pages.append(
                    {
                        "text": text,
                        "page": page_number,
                        "source": source,
                        "total_pages": len(pdf)
                    }
                )

    return pages
